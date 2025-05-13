import streamlit as st
import torch
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
from diffusers import StableDiffusionPipeline
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Streamlit page config
st.set_page_config(page_title="InnovAIte – AI-Powered Prompt Studio", layout="wide", initial_sidebar_state="expanded")

# Model names
TEXT_MODEL = "meta-llama/Llama-2-13b-chat-hf"
IMAGE_MODEL = "stabilityai/stable-diffusion-2-1"

# Load text generation model
@st.cache_resource
def load_text_model():
    bnb_config = BitsAndBytesConfig(
        load_in_4bit=True,
        bnb_4bit_compute_dtype=torch.float16,
        bnb_4bit_use_double_quant=True
    )
    tokenizer = AutoTokenizer.from_pretrained(TEXT_MODEL)
    model = AutoModelForCausalLM.from_pretrained(
        TEXT_MODEL,
        quantization_config=bnb_config,
        device_map="auto"
    )
    return model, tokenizer

# Load image generation model
@st.cache_resource
def load_image_model():
    return StableDiffusionPipeline.from_pretrained(
        IMAGE_MODEL,
        torch_dtype=torch.float16,
        device_map="auto"
    )

# Load models
model, tokenizer = load_text_model()
pipe = load_image_model()

# Sidebar Navigation
st.sidebar.title("Navigation")
pages = ["Home", "Prompt Optimization", "Text Generation", "Image Generation"]
selected_page = st.sidebar.radio("Go to", pages)

if selected_page == "Home":
    st.title("Welcome to InnovAIte")
    st.markdown("### AI-Powered Blog & Image Generation")
    st.markdown("""
**Why Use InnovAIte?**
- 📝 Generate high-quality blogs
- 🎨 Create AI-generated images
- ⏳ Save time with automation
- 📥 Download full blog content

**Steps:**
1. Enter your blog topic
2. Optimize your prompt
3. Generate content
4. Download as DOCX
""")

elif selected_page == "Prompt Optimization":
    st.title("Prompt Optimization")

    st.markdown("### Blog Prompt")
    user_topic = st.text_input("Blog topic:")
    tone = st.selectbox("Tone", ["Formal", "Informal", "Creative", "Persuasive", "Inspirational"])
    audience = st.selectbox("Audience", ["Students", "Professionals", "General", "Experts"])
    temperature = st.slider("Creativity (Temperature)", 0.1, 1.0, 0.7)
    length = st.selectbox("Length", ["Short (300–500 words)", "Medium (500–1000 words)", "Long (1000+ words)"])
    style = st.selectbox("Writing Style", ["Storytelling", "Analytical", "Journalistic", "Technical", "Conversational"])

    if user_topic:
        text_prompt = f"Write a {tone} blog post about {user_topic} for {audience}. Creativity Level: {temperature}, Response Length: {length}, Writing Style: {style}."
        st.markdown("### Generated Text Prompt:")
        st.code(text_prompt)

    st.markdown("### Image Prompt")
    image_subject = st.text_input("Image subject:")
    image_style = st.selectbox("Style", ["Photorealistic", "Digital Art", "Watercolor", "Sketch", "Oil Painting"])
    lighting = st.selectbox("Lighting", ["Natural", "Soft", "Dramatic", "Neon-lit", "Cinematic", "Sunset", "Nighttime"])
    angle = st.selectbox("Camera Angle", ["Close-up", "Wide-angle", "Overhead", "First-person", "Isometric"])
    mood = st.selectbox("Mood", ["Mysterious", "Dreamy", "Futuristic", "Cozy", "Surreal", "Hyper-realistic"])
    colors = st.selectbox("Color Scheme", ["Black & White", "Pastel", "Cyberpunk", "Vivid", "Monochrome"])

    if image_subject:
        image_prompt = f"Generate a {image_style} image of {image_subject} with {lighting} lighting, {angle} angle, evoking a {mood} mood, using a {colors} palette."
        st.markdown("### Generated Image Prompt:")
        st.code(image_prompt)

elif selected_page == "Text Generation":
    st.title("Text Generation")
    text_prompt = st.text_input("Enter prompt for blog generation:")

    if st.button("Generate Text"):
        with st.spinner("Generating..."):
            inputs = tokenizer(text_prompt, return_tensors="pt").to(model.device)
            output = model.generate(**inputs, max_new_tokens=500, temperature=0.7)
            generated_text = tokenizer.decode(output[0], skip_special_tokens=True)
            st.session_state.generated_text = generated_text
            st.markdown("### Generated Blog:")
            st.write(generated_text)

elif selected_page == "Image Generation":
    st.title("Image Generation")
    image_prompt = st.text_input("Enter image prompt:")
    num_images = st.slider("Number of images", 1, 5, 1)

    if st.button("Generate Images"):
        with st.spinner("Generating..."):
            images = [pipe(image_prompt, guidance_scale=7.5, height=768, width=768).images[0] for _ in range(num_images)]
            st.session_state.generated_images = images
            for idx, img in enumerate(images):
                st.image(img, caption=f"Image {idx+1}", use_column_width=True)

# Word document download
if "generated_text" in st.session_state or "generated_images" in st.session_state:
    st.markdown("### Download Generated Content")
    if st.button("Download as DOCX"):
        doc = Document()

        if "generated_text" in st.session_state:
            doc.add_heading("Generated Blog", level=1)
            doc.add_paragraph(st.session_state.generated_text)

        if "generated_images" in st.session_state:
            doc.add_heading("Generated Images", level=1)
            for idx, image in enumerate(st.session_state.generated_images):
                image_io = BytesIO()
                image.save(image_io, format="PNG")
                image_io.seek(0)
                doc.add_paragraph(f"Image {idx+1}:")
                doc.add_picture(image_io, width=Inches(4))

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        st.download_button(
            label="📄 Download Word File",
            data=doc_io,
            file_name="InnovAIte_Content.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
